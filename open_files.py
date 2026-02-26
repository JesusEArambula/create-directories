import docx

def replace_text(filename, output_filename, replacements, dates_replacements):
    # Open docx file
    doc = docx.Document(filename)

    # Loop throuh each paragraph in the docx file
    for p in doc.paragraphs:
        for key in replacements:
        # If the key appears in the text, replace it with the value
            if key in p.text:
                # Replacing text at the paragraph level often loses formatting
                p.text = p.text.replace(key, replacements[key])
                print(f"Replaced '{key}' with '{replacements[key]}'")
    
    # You also need to check tables, headers, and footers separately
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key in replacements:
                        if key in p.text:
                            p.text = p.text.replace(key, replacements[key])
                            print(f"Replaced '{key}' with '{replacements[key]}'")
                    for date_replacement in dates_replacements:
                        if date_replacement in p.text:
                            p.text = p.text.replace(date_replacement, dates_replacements[date_replacement])
                            print(f"Replaced '{date_replacement}' with '{dates_replacements[date_replacement]}'")

    doc.save(output_filename)
    print(f"Document saved as {output_filename} using python-docx-replace.")
